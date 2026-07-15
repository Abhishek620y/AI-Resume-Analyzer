"""
Raw text extraction from resume files.

PDF: pdfplumber (primary) — good text-layout fidelity for resumes.
     Falls back to PyMuPDF (fitz) if pdfplumber returns nothing usable
     (e.g. certain PDF encodings pdfplumber struggles with).
DOCX: python-docx — reads paragraphs and table cells (some resume
     templates put content in tables).
"""
import re

import pdfplumber
import fitz  # PyMuPDF
from docx import Document


def extract_text_from_pdf(file_path: str) -> str:
    text_parts: list[str] = []

    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception:
        text_parts = []

    text = "\n".join(text_parts).strip()
    if text:
        return text

    # Fallback: PyMuPDF
    try:
        doc = fitz.open(file_path)
        text_parts = [page.get_text() for page in doc]
        doc.close()
    except Exception:
        text_parts = []

    return "\n".join(text_parts).strip()


def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    text_parts: list[str] = []

    # Preserve blank paragraphs as empty lines — section/entry splitting
    # relies on blank-line separation between resume entries.
    for para in doc.paragraphs:
        text_parts.append(para.text)

    # Some resume templates place content in tables (e.g. skills grids).
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    text = "\n".join(text_parts)
    # Collapse 3+ consecutive blank lines down to exactly one blank line.
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text(file_path: str, file_extension: str) -> str:
    ext = file_extension.lower().lstrip(".")
    if ext == "pdf":
        return extract_text_from_pdf(file_path)
    elif ext == "docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")
